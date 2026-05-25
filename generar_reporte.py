"""
Genera reporte de diagnóstico de Idelcom.Api en formato .docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Márgenes ──────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

# ── Helpers ───────────────────────────────────────────────────────────────────

AZUL_OSCURO = RGBColor(0x1F, 0x49, 0x7D)
ROJO        = RGBColor(0xC0, 0x00, 0x00)
NARANJA     = RGBColor(0xE3, 0x6C, 0x09)
VERDE       = RGBColor(0x37, 0x86, 0x35)
NEGRO       = RGBColor(0x00, 0x00, 0x00)

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, sides=("top","bottom","left","right"), color="CCCCCC", size="4"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(15)
    run.font.color.rgb = AZUL_OSCURO
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    btm = OxmlElement("w:bottom")
    btm.set(qn("w:val"),   "single")
    btm.set(qn("w:sz"),    "6")
    btm.set(qn("w:space"), "1")
    btm.set(qn("w:color"), "1F497D")
    pBdr.append(btm)
    pPr.append(pBdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(12)
    run.font.color.rgb = AZUL_OSCURO
    return p

def body(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.bold   = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    if bold_prefix is not None:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(text)
        r2.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr  = p._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F2F2F2")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    return p

def issue_table(rows):
    COL_W = [Inches(1.1), Inches(2.6), Inches(2.8)]
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = tbl.rows[0].cells
    for i, (txt, w) in enumerate(zip(["Severidad","Archivo / Ubicacion","Problema"], COL_W)):
        hdr[i].width = w
        set_cell_bg(hdr[i], "1F497D")
        set_cell_border(hdr[i], color="FFFFFF")
        p = hdr[i].paragraphs[0]
        run = p.add_run(txt)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
    COLOR_SEV = {
        "CRITICO": ("C00000", "FFEDED"),
        "ALTO":    ("E36C09", "FFF4EC"),
        "MEDIO":   ("7F6000", "FFFBE6"),
        "BAJO":    ("375623", "F0F7EE"),
    }
    for i, (sev, loc, desc) in enumerate(rows):
        row = tbl.add_row()
        fg  = COLOR_SEV.get(sev, ("000000", "FFFFFF"))[0]
        bg  = COLOR_SEV.get(sev, ("000000", "FFFFFF"))[1]
        for c in row.cells:
            set_cell_bg(c, "FFFFFF" if i % 2 == 0 else "F8F8F8")
            set_cell_border(c, color="DDDDDD", size="2")
        row.cells[0].width = COL_W[0]
        row.cells[1].width = COL_W[1]
        row.cells[2].width = COL_W[2]
        set_cell_bg(row.cells[0], bg)
        r0 = row.cells[0].paragraphs[0].add_run(sev)
        r0.bold = True
        r0.font.size = Pt(9)
        r0.font.color.rgb = RGBColor.from_string(fg)
        row.cells[1].paragraphs[0].add_run(loc).font.size = Pt(9)
        row.cells[2].paragraphs[0].add_run(desc).font.size = Pt(9)
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
#  PORTADA
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run("IDELCOM ERP  GlueMark API")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = AZUL_OSCURO

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Diagnostico de Estabilidad y Seguridad en Produccion")
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f"Fecha: {datetime.date.today().strftime('%d de %B de %Y')}")
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  1. RESUMEN EJECUTIVO
# ══════════════════════════════════════════════════════════════════════════════
heading1("1. Resumen Ejecutivo")

body(
    "Se revisaron los 2.669 archivos .cs del proyecto Idelcom.Api (solucion GlueMark) "
    "con foco en estabilidad, seguridad y calidad de codigo. La revision incluyo: "
    "configuracion del servidor, workers en segundo plano, repositorios de datos, "
    "middleware, seguridad JWT, manejo de memoria y cadenas de conexion a SQL Server.",
)
body(
    "Se encontraron 4 problemas CRITICOS que explican directamente las caidas "
    "intermitentes reportadas en produccion (ERR_CONNECTION_TIMED_OUT en "
    "/api/auth/session), y 9 problemas adicionales de seguridad y calidad "
    "que deben corregirse antes del proximo despliegue.",
)

tbl = doc.add_table(rows=2, cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
cabeceras = ["CRITICO", "ALTO", "MEDIO", "BAJO"]
valores   = ["4", "4", "3", "2"]
COLORES   = ["C00000", "E36C09", "7F6000", "375623"]
FONDOS    = ["FFEDED", "FFF4EC", "FFFBE6", "F0F7EE"]
for i, (cab, val, fg, bg) in enumerate(zip(cabeceras, valores, COLORES, FONDOS)):
    c_h = tbl.rows[0].cells[i]
    c_v = tbl.rows[1].cells[i]
    set_cell_bg(c_h, fg)
    set_cell_bg(c_v, bg)
    set_cell_border(c_h, color="FFFFFF")
    set_cell_border(c_v, color="DDDDDD", size="2")
    r = c_h.paragraphs[0].add_run(cab)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    c_h.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = c_v.paragraphs[0].add_run(val)
    r2.bold = True
    r2.font.size = Pt(20)
    r2.font.color.rgb = RGBColor.from_string(fg)
    c_v.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  2. DIAGNOSTICO DE LAS CAIDAS EN PRODUCCION
# ══════════════════════════════════════════════════════════════════════════════
heading1("2. Diagnostico de las Caidas en Produccion")

body(
    "El sintoma observado es: el sistema funciona correctamente durante un periodo "
    "y luego el frontend muestra 'No pudimos conectar con el sistema' con "
    "ERR_CONNECTION_TIMED_OUT en /api/auth/session. Limpiar la cache del navegador "
    "restaura el acceso temporalmente. Usuarios nuevos no pueden entrar durante "
    "el periodo de caida.",
)
body("Hay dos causas independientes que se combinan para producir este comportamiento:")

heading2("2.1  Causa Raiz: Saturacion de SQL Server Express")

body(
    "La base de datos de produccion corre sobre SQL Server Express "
    "(SRV-ERP-SAPI\\SQLEXPRESS), edicion con las siguientes restricciones de hardware:",
)
bullet("Maximo 1 GB de RAM para el motor de base de datos.")
bullet("Maximo 1 CPU (1 nucleo logico).")
bullet("Tamano maximo de base de datos: 10 GB.")
body(
    "El proyecto registra 4 Hosted Services (workers en segundo plano) que "
    "ejecutan consultas a SQL Server de forma continua:"
)

issue_table([
    ("CRITICO", "NotificationsOutboxWorker.cs",  "Polling cada 2 segundos. Sin delay cuando hay items pendientes. Puede generar rafagas de decenas de consultas por segundo."),
    ("CRITICO", "EmailOutboxWorker.cs",           "Polling cada 10 segundos, 20 correos por lote. Usa varios SPs por correo."),
    ("CRITICO", "OpportunityAlertsWorker.cs",     "Ejecuta 3 Stored Procedures complejos cada 60 segundos. @BUSINESS_ID hardcodeado = 1."),
    ("ALTO",    "SupportAlertsWorker.cs",          "Ejecuta 1 SP durante la ventana 8am - 9am cada 5 minutos."),
])

body(
    "Bajo carga normal de usuarios mas los 4 workers, SQL Server Express supera "
    "su limite de RAM. El motor responde lentamente o deja de responder, "
    "lo que provoca que las peticiones HTTP agoten su timeout TCP (ERR_CONNECTION_TIMED_OUT)."
)

heading2("2.2  Sin CommandTimeout en los Repositorios")

body(
    "Ninguno de los 142 repositorios configura CommandTimeout en los SqlCommand. "
    "El valor por defecto de SQL Server es 30 segundos."
)
body(
    "Cuando SQL Express se satura, las queries pueden tardar mas de 30 segundos. "
    "Los threads del thread pool de .NET quedan bloqueados esperando respuesta. "
    "Con suficientes queries lentas en paralelo, el thread pool se agota y "
    "el servidor deja de aceptar nuevas conexiones TCP."
)
code_block(
    "// Estado actual (sin timeout):\n"
    "using var cmd = new SqlCommand(\"SP_WS_...\", cn)\n"
    "{\n"
    "    CommandType = CommandType.StoredProcedure\n"
    "};\n\n"
    "// Correccion (agregar timeout):\n"
    "using var cmd = new SqlCommand(\"SP_WS_...\", cn)\n"
    "{\n"
    "    CommandType = CommandType.StoredProcedure,\n"
    "    CommandTimeout = 15  // maximo 15 segundos\n"
    "};"
)

heading2("2.3  Por que Limpiar Cache lo Soluciona")

body(
    "El frontend tiene un interceptor de red que, al recibir un error de conexion, "
    "persiste el estado 'backend caido' en el localStorage del navegador. "
    "Una vez marcado como caido, bloquea todas las peticiones sin reintentar."
)
bullet("Usuarios existentes: limpiar cache borra el flag del localStorage, el frontend reintenta, el backend ya se recupero, funciona.")
bullet("Usuarios nuevos: si acceden mientras el backend esta caido, el flag se graba en su primera visita y nunca pueden entrar.")
body("Fix requerido en el frontend:", bold=True)
code_block(
    "// Agregar retry automatico en el interceptor del frontend:\n"
    "setInterval(async () => {\n"
    "  if (backendDown) {\n"
    "    const ok = await fetch('/api/Health', { method: 'HEAD' })\n"
    "      .then(r => r.ok).catch(() => false);\n"
    "    if (ok) { backendDown = false; window.location.reload(); }\n"
    "  }\n"
    "}, 30_000);"
)

# ══════════════════════════════════════════════════════════════════════════════
#  3. HALLAZGOS DE SEGURIDAD
# ══════════════════════════════════════════════════════════════════════════════
heading1("3. Hallazgos de Seguridad")

heading2("3.1  Secretos Expuestos en Codigo Fuente")

body(
    "Los archivos appsettings.json contienen credenciales sensibles que estan "
    "en el repositorio de codigo. Cualquier persona con acceso al repo puede "
    "comprometer el sistema completo:"
)

issue_table([
    ("CRITICO", "appsettings.json: JwtSetting:Key",      "Clave de firma JWT expuesta. Permite fabricar tokens validos para cualquier usuario."),
    ("CRITICO", "appsettings.json: ConnectionStrings",   "Password de SQL Server (usuario sa) expuesta. Acceso total a la base de datos."),
    ("CRITICO", "appsettings.json: Smtp:Pass",           "Password del servidor SMTP expuesto."),
    ("ALTO",    "appsettings.json: ApiPeru:ApiKey",       "API key de servicio externo expuesta."),
])

body("Solucion: mover secretos a variables de entorno del servidor.", bold=True)
code_block(
    "# En el servidor Windows, configurar variables de entorno:\n"
    "# ConnectionStrings__connection=Server=...;Password=SecurePass;\n"
    "# JwtSetting__Key=ClaveSegura64CaracteresRandom\n"
    "# Smtp__Pass=PasswordSMTP\n\n"
    "# En appsettings.Production.json solo dejar config sin secretos:\n"
    '{\n'
    '  "JwtSetting": {\n'
    '    "Issuer": "https://api.idelcom.pe",\n'
    '    "Audience": "https://api.idelcom.pe"\n'
    '  }\n'
    '}'
)

heading2("3.2  Swagger Habilitado en Produccion")

body(
    'appsettings.json tiene "EnableSwaggerInProduction": true. Esto expone toda '
    "la documentacion de la API publicamente en https://api.idelcom.pe/swagger, "
    "facilitando ataques dirigidos."
)
code_block(
    '// appsettings.Production.json\n'
    '{\n'
    '  "EnableSwaggerInProduction": false\n'
    '}'
)

heading2("3.3  RequireHttpsMetadata = false en JWT")

body(
    "En Program.cs, el middleware JWT tiene RequireHttpsMetadata = false, "
    "permitiendo que tokens se transmitan por HTTP sin cifrar. "
    "Cambiar a: config.RequireHttpsMetadata = !builder.Environment.IsDevelopment();"
)

heading2("3.4  CORS con AllowedOrigins: [*] en archivo base")

body(
    'El archivo appsettings.json base tiene AllowedOrigins: ["*"], '
    "abriendo CORS a cualquier origen como valor por defecto. "
    "Eliminar el wildcard del archivo base y definir origenes especificos "
    "en cada appsettings.[Env].json."
)

# ══════════════════════════════════════════════════════════════════════════════
#  4. OTROS HALLAZGOS
# ══════════════════════════════════════════════════════════════════════════════
heading1("4. Otros Hallazgos de Calidad")

heading2("4.1  Bug: OpportunityAlertsWorker con BUSINESS_ID = 1 fijo")

body(
    "En OpportunityAlertsWorker.cs (lineas 49 y 57), los SPs "
    "SP_WS_PROCESS_INACTIVITY_ALERTS y SP_WS_PROCESS_PRESALES_ALERTS reciben "
    "@BUSINESS_ID = 1 hardcodeado. Si la empresa tiene ID distinto de 1, "
    "o si se agregan nuevas empresas, las alertas nunca se procesaran para ellas."
)
code_block(
    "// Actual (incorrecto):\n"
    'cmdInactivity.Parameters.AddWithValue("@BUSINESS_ID", 1);\n\n'
    "// Correccion: leer de configuracion o iterar empresas activas\n"
    'var businessId = _configuration.GetValue<int>("AppSettings:BusinessId");\n'
    'cmdInactivity.Parameters.AddWithValue("@BUSINESS_ID", businessId);'
)

heading2("4.2  Console.WriteLine en Produccion")

body(
    "Se encontraron 6 llamadas a Console.WriteLine en archivos de produccion. "
    "Generan bloqueos de I/O bajo carga y no quedan en el sistema de logs. "
    "Reemplazar por _logger.LogInformation / _logger.LogError:"
)
issue_table([
    ("MEDIO", "EmailOutboxWorker.cs:67",             "Console.WriteLine procesando item"),
    ("MEDIO", "SalesQuotationRepository.cs:96, 161", "Console.WriteLine de excepciones SQL"),
    ("MEDIO", "EmailTemplateRenderer.cs:38",          "Console.WriteLine de recursos embebidos"),
    ("MEDIO", "NotificationsHub.cs:13",               "Console.WriteLine de conexiones SignalR"),
])

heading2("4.3  NotificationsOutboxWorker Sin Delay en Happy Path")

body(
    "Cuando hay notificaciones pendientes, el worker procesa el lote y reinicia "
    "el ciclo inmediatamente sin pausa. Con muchas notificaciones en cola, "
    "genera carga continua sobre la BD. Agregar un delay minimo de 500 ms:"
)
code_block(
    "// Agregar al final del bloque try, despues de MarkProcessedAsync:\n"
    "await Task.Delay(500, stoppingToken);"
)

# ══════════════════════════════════════════════════════════════════════════════
#  5. PLAN DE ACCION
# ══════════════════════════════════════════════════════════════════════════════
heading1("5. Plan de Accion Priorizado")

tbl2 = doc.add_table(rows=1, cols=4)
tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, txt in enumerate(["#", "Accion", "Esfuerzo", "Impacto"]):
    c = tbl2.rows[0].cells[i]
    set_cell_bg(c, "1F497D")
    set_cell_border(c, color="FFFFFF")
    r = c.paragraphs[0].add_run(txt)
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    r.font.size = Pt(9)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

acciones = [
    ("1", "Agregar CommandTimeout=15 en los 142 repositorios. Usar busqueda/reemplazo global en Visual Studio.", "Alto", "CRITICO"),
    ("2", "Reducir polling de NotificationsOutboxWorker de 2s a 10s. Agregar delay 500ms en happy path.", "Bajo", "CRITICO"),
    ("3", "Agregar limites al connection pool en la cadena de conexion: Max Pool Size=30; Connection Timeout=10;", "Bajo", "CRITICO"),
    ("4", "Evaluar upgrade de SQL Server Express a Developer/Standard Edition en el servidor.", "Medio", "CRITICO"),
    ("5", "Mover secretos (JWT Key, SQL password, SMTP password) a variables de entorno del servidor.", "Medio", "CRITICO"),
    ("6", "Deshabilitar Swagger: EnableSwaggerInProduction: false en appsettings.Production.json", "Bajo", "ALTO"),
    ("7", "Corregir @BUSINESS_ID hardcodeado en OpportunityAlertsWorker.", "Bajo", "ALTO"),
    ("8", "Cambiar RequireHttpsMetadata a !builder.Environment.IsDevelopment()", "Bajo", "ALTO"),
    ("9", "Reemplazar los 6 Console.WriteLine por _logger en archivos de produccion.", "Bajo", "MEDIO"),
    ("10","Eliminar AllowedOrigins: ['*'] del appsettings.json base.", "Bajo", "MEDIO"),
    ("11","Fix del interceptor de red del frontend: agregar retry automatico cada 30s.", "Medio", "CRITICO"),
]

IMP_COLOR = {"CRITICO": "C00000", "ALTO": "E36C09", "MEDIO": "7F6000"}
for i, (num, accion, esfuerzo, impacto) in enumerate(acciones):
    row = tbl2.add_row()
    bg = "FFFFFF" if i % 2 == 0 else "F5F5F5"
    for c in row.cells:
        set_cell_bg(c, bg)
        set_cell_border(c, color="DDDDDD", size="2")
    row.cells[0].paragraphs[0].add_run(num).font.size = Pt(9)
    row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row.cells[1].paragraphs[0].add_run(accion).font.size = Pt(9)
    row.cells[2].paragraphs[0].add_run(esfuerzo).font.size = Pt(9)
    row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = row.cells[3].paragraphs[0].add_run(impacto)
    r3.font.size = Pt(9)
    r3.bold = True
    r3.font.color.rgb = RGBColor.from_string(IMP_COLOR.get(impacto, "000000"))
    row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  6. FIX INMEDIATO
# ══════════════════════════════════════════════════════════════════════════════
heading1("6. Fix Inmediato para Estabilizar Produccion Hoy")

body("Sin tocar codigo, estas acciones en el servidor mejoran la estabilidad en minutos:", bold=True)

body("Paso 1 - IIS Application Pool:", bold=True)
body("Abrir IIS Manager > Application Pools > Pool de GlueMark > Advanced Settings:")
code_block(
    "Idle Time-out (minutes): 0\n"
    "Regular Time Interval (minutes): 0\n"
    "Private Memory Limit (KB): 0"
)

body("Paso 2 - Cadena de conexion en appsettings.Production.json:", bold=True)
code_block(
    '"connection": "Server=SRV-ERP-SAPI\\\\SQLEXPRESS;\n'
    '    DataBase=IDELCOM; User ID=sa; Password=...;\n'
    '    TrustServerCertificate=True;\n'
    '    Max Pool Size=30;\n'
    '    Connection Timeout=10;"'
)

body("Paso 3 - Reducir polling de NotificationsOutboxWorker:", bold=True)
code_block(
    "// Infrastructure/Notifications/NotificationsOutboxWorker.cs linea 47:\n"
    "// Cambiar:\n"
    "await Task.Delay(2000, stoppingToken);\n"
    "// A:\n"
    "await Task.Delay(10_000, stoppingToken);"
)

# ══════════════════════════════════════════════════════════════════════════════
#  7. ASPECTOS POSITIVOS
# ══════════════════════════════════════════════════════════════════════════════
heading1("7. Aspectos Positivos del Proyecto")

body("El proyecto tiene una arquitectura solida con buenas practicas en varias areas:")
bullet("Clean Architecture bien implementada con separacion clara entre capas.")
bullet("Manejo de conexiones SQL correcto: todos los SqlConnection y SqlCommand usan using, sin fugas de conexiones.")
bullet("Servicios de seguridad bien disenados: TokenBlacklist, LoginAttemptService y RefreshTokenService usan IDistributedCache con TTL.")
bullet("Background workers con IServiceScopeFactory: los workers crean scopes correctamente para servicios scoped.")
bullet("ExceptionHandlingMiddleware centralizado: todos los errores se manejan en un solo punto sin exponer stack traces.")
bullet("Patron Outbox para emails y notificaciones: buena practica para garantizar entrega eventual.")
bullet("HealthController existente en /api/Health: base solida para monitoreo.")

doc.add_paragraph()
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
run = p.add_run("-- Fin del Reporte --")
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── Guardar ───────────────────────────────────────────────────────────────────
OUT = "/sessions/elegant-lucid-galileo/mnt/Idelcom.Api/Diagnostico_Produccion_Idelcom.docx"
doc.save(OUT)
print(f"Guardado: {OUT}")
